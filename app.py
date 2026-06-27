"""
===============================================================================
EXIM ENTERPRISE RECONCILIATION & ECUS GENERATOR
===============================================================================
Version  : 4.0.0 (Production-Hardened)
Languages: Vietnamese · English · Chinese
Formats  : Excel (xlsx/xls/xlsb/xlsm) · CSV · TXT · PDF · JPG/PNG (OCR)
Fixes    : All KeyError bugs, Chinese locale added, status priority logic
           corrected, missing ci_* i18n keys added, pdfplumber hoisted to
           top-level import, status colour-priority (Red > Yellow > Green).
===============================================================================
"""
from __future__ import annotations

import io
import os
import re
import traceback
import warnings
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
import pdfplumber
import streamlit as st
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

warnings.filterwarnings("ignore")


# =============================================================================
# 1.  LOGGING
# =============================================================================

class LogLevel(Enum):
    INFO     = "INFO"
    WARNING  = "WARNING"
    ERROR    = "ERROR"

@dataclass
class SystemLog:
    timestamp : datetime
    level     : LogLevel
    message   : str
    details   : Optional[str] = None

class SystemLogger:
    _KEY = "sys_logs"

    def __init__(self):
        if self._KEY not in st.session_state:
            st.session_state[self._KEY] = []

    def _push(self, level: LogLevel, msg: str, details: str = None):
        st.session_state[self._KEY].append(
            SystemLog(datetime.now(), level, msg, details)
        )

    def info   (self, m: str):             self._push(LogLevel.INFO,    m)
    def warning(self, m: str):             self._push(LogLevel.WARNING, m)
    def error  (self, m: str, d: str = None): self._push(LogLevel.ERROR, m, d)

    def all(self) -> List[SystemLog]:
        return st.session_state.get(self._KEY, [])

    def clear(self):
        st.session_state[self._KEY] = []


log = SystemLogger()


# =============================================================================
# 2.  i18n  (VI · EN · ZH — every key present in all three locales)
# =============================================================================

# Column-index label helpers (same text can be shared across langs)
_CI_LABELS = {
    "vi": {
        "ci_inv_mat": "Invoice – cột Mã vật tư (index):",
        "ci_inv_qty": "Invoice – cột Số lượng (index):",
        "ci_mb_mat":  "MB52 – cột Mã vật tư (index):",
        "ci_mb_stk":  "MB52 – cột Tồn kho (index):",
        "ci_iop_mat": "IOP01 – cột Mã vật tư (index):",
        "ci_iop_rt":  "IOP01 – cột Tỷ lệ NL (index):",
        "ci_iop_ut":  "IOP01 – cột ĐVT (index):",
        "ci_iop_hs":  "IOP01 – cột Mã HS (index):",
        "ci_hd_mat":  "HD03 – cột Mã NL (index):",
        "ci_hd_tk":   "HD03 – cột Số TK (index):",
        "ci_hd_bal":  "HD03 – cột Tồn (index):",
        "ci_hd_pri":  "HD03 – cột Đơn giá (index):",
        "ci_hd_dsc":  "HD03 – cột Tên hàng (index):",
        "ci_zmm_mat": "ZMM12 – cột Mã vật tư (index):",
        "ci_zmm_qty": "ZMM12 – cột Số lượng (index):",
    },
    "en": {
        "ci_inv_mat": "Invoice – Material col (index):",
        "ci_inv_qty": "Invoice – Qty col (index):",
        "ci_mb_mat":  "MB52 – Material col (index):",
        "ci_mb_stk":  "MB52 – Stock col (index):",
        "ci_iop_mat": "IOP01 – Material col (index):",
        "ci_iop_rt":  "IOP01 – Rate col (index):",
        "ci_iop_ut":  "IOP01 – UOM col (index):",
        "ci_iop_hs":  "IOP01 – HS Code col (index):",
        "ci_hd_mat":  "HD03 – Material col (index):",
        "ci_hd_tk":   "HD03 – Decl. No col (index):",
        "ci_hd_bal":  "HD03 – Balance col (index):",
        "ci_hd_pri":  "HD03 – Price col (index):",
        "ci_hd_dsc":  "HD03 – Description col (index):",
        "ci_zmm_mat": "ZMM12 – Material col (index):",
        "ci_zmm_qty": "ZMM12 – Qty col (index):",
    },
    "zh": {
        "ci_inv_mat": "发票 – 物料列 (索引):",
        "ci_inv_qty": "发票 – 数量列 (索引):",
        "ci_mb_mat":  "MB52 – 物料列 (索引):",
        "ci_mb_stk":  "MB52 – 库存列 (索引):",
        "ci_iop_mat": "IOP01 – 物料列 (索引):",
        "ci_iop_rt":  "IOP01 – 换算率列 (索引):",
        "ci_iop_ut":  "IOP01 – 计量单位列 (索引):",
        "ci_iop_hs":  "IOP01 – HS编码列 (索引):",
        "ci_hd_mat":  "HD03 – 物料列 (索引):",
        "ci_hd_tk":   "HD03 – 报关单列 (索引):",
        "ci_hd_bal":  "HD03 – 余量列 (索引):",
        "ci_hd_pri":  "HD03 – 单价列 (索引):",
        "ci_hd_dsc":  "HD03 – 描述列 (索引):",
        "ci_zmm_mat": "ZMM12 – 物料列 (索引):",
        "ci_zmm_qty": "ZMM12 – 数量列 (索引):",
    },
}

LANG: Dict[str, Dict] = {
    # ─────────────────────────── VIETNAMESE ───────────────────────────────
    "🇻🇳 Tiếng Việt": {
        **_CI_LABELS["vi"],
        "gate_title"        : "🌐 HỆ THỐNG ĐỐI SOÁT CHỨNG TỪ SCM EXIM",
        "gate_sub"          : "Vui lòng đăng nhập để truy cập công cụ tự động hóa",
        "pwd_label"         : "Mã bảo mật hệ thống:",
        "pwd_err"           : "❌ Mã truy cập không hợp lệ. Vui lòng thử lại!",
        "main_title"        : "🏢 E23-E54: ĐỐI SOÁT & SINH CHUỖI ECUS TỰ ĐỘNG",
        "main_sub"          : "Động cơ FIFO Allocation: Gom nhóm ➔ Kiểm tồn kho ➔ Chẻ dòng tự động.",
        "tab_engine"        : "⚙️ ENGINE CHÍNH",
        "tab_dashboard"     : "📊 DASHBOARD",
        "tab_logs"          : "📝 NHẬT KÝ HỆ THỐNG",
        "dl_title"          : "🗂️ BƯỚC 1: NẠP HỒ DỮ LIỆU ĐỐI CHIẾU",
        "dl_sub"            : "Kéo thả 6 file thô từ hệ thống. Hỗ trợ Excel, CSV, PDF, Ảnh scan.",
        "f_inv"             : "1. INVOICE KHAI BÁO (MC)",
        "f_pkl"             : "2. PACKING LIST (MC)",
        "f_hd03"            : "3. SỔ HD03 (HẢI QUAN)",
        "f_zmm12"           : "4. SAP ZMM12 (Kế toán)",
        "f_iop01"           : "5. SAP IOP01 (Quy đổi)",
        "f_mb52"            : "6. SAP MB52/MB51 (Kho)",
        "eng_title"         : "🧠 BƯỚC 2: ENGINE KIỂM TOÁN & CHẺ DÒNG FIFO",
        "spin_msg"          : "Cỗ máy đang phân tích ma trận và chẻ dòng FIFO...",
        "succ_msg"          : "🎉 Hoàn tất! Tiền kiểm và phân bổ FIFO đã thành công.",
        "err_toggle"        : "🚨 Chỉ hiện dòng có LỖI / CẢNH BÁO",
        "btn_xlsx"          : "📥 TẢI FILE EXCEL ECUS",
        "btn_docx"          : "📝 TẢI BIÊN BẢN WORD",
        "miss_files"        : "💡 Vui lòng nạp đủ 6 file chứng từ để khởi động động cơ ECUS.",
        "file_broken"       : "🚨 File **{name}** không đọc được. Kiểm tra Log để biết chi tiết.",
        "st_err_pkl"        : "🔴 LỆCH SỐ LƯỢNG INV({inv}) ≠ PKL({pkl}) | ",
        "st_err_mb52"       : "🔴 ÂM KHO VẬT LÝ MB52 (Tồn thực: {stock})",
        "st_err_hd03_miss"  : "🔴 MÃ KHÔNG TỒN TẠI TRONG SỔ HD03",
        "st_err_hd03_empty" : "🔴 SỔ HD03 ĐÃ BỊ TRỪ HẾT TỒN",
        "st_warn_zmm"       : "🟡 TỒN KẾ TOÁN ZMM12 KHÔNG ĐỦ | ",
        "st_ok"             : "🟢 HỢP LỆ – SẴN SÀNG KHAI BÁO",
        "st_warn_chk"       : "CẦN RÀ SOÁT SỔ KẾ TOÁN",
        "m_total"           : "Tổng Dòng",
        "m_ok"              : "✅ Hợp lệ",
        "m_warn"            : "🟡 Cảnh báo",
        "m_err"             : "🔴 Lỗi",
        "m_value"           : "💰 Tổng Trị giá (USD)",
        "dash_title"        : "Phân tích Sức khỏe Dữ liệu",
        "dash_health"       : "Phân bổ Trạng thái",
        "dash_top5"         : "Top 5 Mã hàng theo Trị giá",
        "dash_empty"        : "Chạy Engine trước để vẽ biểu đồ.",
        "log_title"         : "📝 Nhật ký hệ thống",
        "log_empty"         : "Chưa có log.",
        "log_no_detail"     : "Không có chi tiết kỹ thuật.",
        "log_clear"         : "🗑️ Xóa nhật ký",
        "tol_w"             : "Dung sai Vải/Cân ký (KGM, MTK):",
        "tol_c"             : "Dung sai Đếm chiếc (PCE, PRS):",
        "cfg_tol"           : "⚖️ THIẾT LẬP DUNG SAI",
        "cfg_col"           : "🔎 TÙY CHỈNH CỘT INDEX",
        "c_mshq"            : "Mã số HH (HS)",
        "c_malieu"          : "Mã Nguyên Liệu",
        "c_ten"             : "Tên Hàng",
        "c_dvt"             : "ĐVT",
        "c_luong"           : "Lượng Khai Báo",
        "c_dgia"            : "Đơn Giá",
        "c_tgia"            : "Trị Giá",
        "c_tk"              : "Số TK Gốc",
        "c_ref"             : "Mã Tham Chiếu",
        "c_stat"            : "TRẠNG THÁI",
        "logout"            : "🚪 Đăng xuất",
    },

    # ─────────────────────────── ENGLISH ──────────────────────────────────
    "🇺🇸 English": {
        **_CI_LABELS["en"],
        "gate_title"        : "🌐 SCM EXIM RECONCILIATION SYSTEM",
        "gate_sub"          : "Please log in to access the automation tool",
        "pwd_label"         : "System Security Code:",
        "pwd_err"           : "❌ Invalid access code. Please try again!",
        "main_title"        : "🏢 E23-E54: AUTO RECONCILIATION & ECUS GENERATOR",
        "main_sub"          : "FIFO Engine: Grouping ➔ Inventory Check ➔ Auto Line-Splitting.",
        "tab_engine"        : "⚙️ MAIN ENGINE",
        "tab_dashboard"     : "📊 DASHBOARD",
        "tab_logs"          : "📝 SYSTEM LOGS",
        "dl_title"          : "🗂️ STEP 1: LOAD RECONCILIATION DATA",
        "dl_sub"            : "Drag & drop 6 raw system files. Excel, CSV, PDF, Scanned Images all supported.",
        "f_inv"             : "1. OFFICIAL INVOICE (MC)",
        "f_pkl"             : "2. PACKING LIST (MC)",
        "f_hd03"            : "3. HD03 LEDGER (CUSTOMS)",
        "f_zmm12"           : "4. SAP ZMM12 (Accounting)",
        "f_iop01"           : "5. SAP IOP01 (Conversion)",
        "f_mb52"            : "6. SAP MB52/MB51 (Physical Stock)",
        "eng_title"         : "🧠 STEP 2: AUDIT ENGINE & FIFO LINE-SPLITTING",
        "spin_msg"          : "Engine is analyzing matrix and executing FIFO allocation...",
        "succ_msg"          : "🎉 Done! Pre-check and FIFO allocation succeeded.",
        "err_toggle"        : "🚨 Show only ERROR / WARNING rows",
        "btn_xlsx"          : "📥 DOWNLOAD ECUS EXCEL",
        "btn_docx"          : "📝 DOWNLOAD WORD MEMO",
        "miss_files"        : "💡 Please upload all 6 documents to start the ECUS engine.",
        "file_broken"       : "🚨 File **{name}** could not be parsed. See Logs for details.",
        "st_err_pkl"        : "🔴 QTY MISMATCH INV({inv}) ≠ PKL({pkl}) | ",
        "st_err_mb52"       : "🔴 NEGATIVE PHYSICAL STOCK MB52 (Avail: {stock})",
        "st_err_hd03_miss"  : "🔴 ITEM NOT FOUND IN HD03 LEDGER",
        "st_err_hd03_empty" : "🔴 HD03 LEDGER FULLY DEPLETED",
        "st_warn_zmm"       : "🟡 ZMM12 ACCOUNTING STOCK INSUFFICIENT | ",
        "st_ok"             : "🟢 VALID – READY TO DECLARE",
        "st_warn_chk"       : "REVIEW ACCOUNTING LEDGER",
        "m_total"           : "Total Rows",
        "m_ok"              : "✅ Valid",
        "m_warn"            : "🟡 Warnings",
        "m_err"             : "🔴 Errors",
        "m_value"           : "💰 Total Value (USD)",
        "dash_title"        : "Data Health Analytics",
        "dash_health"       : "Status Distribution",
        "dash_top5"         : "Top 5 Items by Value",
        "dash_empty"        : "Run the Engine first to populate charts.",
        "log_title"         : "📝 System Logs",
        "log_empty"         : "No logs recorded yet.",
        "log_no_detail"     : "No technical details available.",
        "log_clear"         : "🗑️ Clear Logs",
        "tol_w"             : "Weight/Fabric tolerance (KGM, MTK):",
        "tol_c"             : "Count tolerance (PCE, PRS):",
        "cfg_tol"           : "⚖️ TOLERANCE SETTINGS",
        "cfg_col"           : "🔎 COLUMN INDEX OVERRIDES",
        "c_mshq"            : "HS Code",
        "c_malieu"          : "Material Code",
        "c_ten"             : "Description",
        "c_dvt"             : "UOM",
        "c_luong"           : "Declared Qty",
        "c_dgia"            : "Unit Price",
        "c_tgia"            : "Total Value",
        "c_tk"              : "Source Decl. No.",
        "c_ref"             : "Ref Code",
        "c_stat"            : "AUDIT STATUS",
        "logout"            : "🚪 Logout",
    },

    # ─────────────────────────── CHINESE ──────────────────────────────────
    "🇨🇳 中文": {
        **_CI_LABELS["zh"],
        "gate_title"        : "🌐 SCM 进出口单证核对系统",
        "gate_sub"          : "请登录以访问自动化工具",
        "pwd_label"         : "系统安全密码:",
        "pwd_err"           : "❌ 访问密码错误，请重试！",
        "main_title"        : "🏢 E23-E54: 自动核对 & ECUS 生成系统",
        "main_sub"          : "FIFO 引擎: 分组 ➔ 库存检查 ➔ 自动行拆分。",
        "tab_engine"        : "⚙️ 主引擎",
        "tab_dashboard"     : "📊 数据看板",
        "tab_logs"          : "📝 系统日志",
        "dl_title"          : "🗂️ 第一步: 加载核对数据",
        "dl_sub"            : "将6个原始系统文件拖放到此处。支持 Excel、CSV、PDF、扫描图片。",
        "f_inv"             : "1. 官方发票 (MC)",
        "f_pkl"             : "2. 装箱单 (MC)",
        "f_hd03"            : "3. HD03 账册 (海关)",
        "f_zmm12"           : "4. SAP ZMM12 (会计)",
        "f_iop01"           : "5. SAP IOP01 (换算)",
        "f_mb52"            : "6. SAP MB52/MB51 (实物库存)",
        "eng_title"         : "🧠 第二步: 审计引擎 & FIFO 行拆分",
        "spin_msg"          : "引擎正在分析矩阵并执行 FIFO 分配...",
        "succ_msg"          : "🎉 完成！预检和 FIFO 分配成功。",
        "err_toggle"        : "🚨 仅显示错误 / 警告行",
        "btn_xlsx"          : "📥 下载 ECUS Excel",
        "btn_docx"          : "📝 下载 Word 备忘录",
        "miss_files"        : "💡 请上传全部 6 个文件以启动 ECUS 引擎。",
        "file_broken"       : "🚨 文件 **{name}** 无法解析，请查看日志了解详情。",
        "st_err_pkl"        : "🔴 数量不符 INV({inv}) ≠ PKL({pkl}) | ",
        "st_err_mb52"       : "🔴 MB52 实物库存为负 (可用: {stock})",
        "st_err_hd03_miss"  : "🔴 该物料在 HD03 账册中不存在",
        "st_err_hd03_empty" : "🔴 HD03 账册余量已全部耗尽",
        "st_warn_zmm"       : "🟡 ZMM12 会计库存不足 | ",
        "st_ok"             : "🟢 有效 – 可申报",
        "st_warn_chk"       : "请核查会计账册",
        "m_total"           : "总行数",
        "m_ok"              : "✅ 有效",
        "m_warn"            : "🟡 警告",
        "m_err"             : "🔴 错误",
        "m_value"           : "💰 总价值 (USD)",
        "dash_title"        : "数据健康分析",
        "dash_health"       : "状态分布",
        "dash_top5"         : "按价值排名前5项",
        "dash_empty"        : "请先运行引擎以填充图表。",
        "log_title"         : "📝 系统日志",
        "log_empty"         : "暂无日志记录。",
        "log_no_detail"     : "无技术详情。",
        "log_clear"         : "🗑️ 清除日志",
        "tol_w"             : "重量/面料容差 (KGM, MTK):",
        "tol_c"             : "计件容差 (PCE, PRS):",
        "cfg_tol"           : "⚖️ 容差设置",
        "cfg_col"           : "🔎 列索引覆盖",
        "c_mshq"            : "海关编码 (HS)",
        "c_malieu"          : "物料代码",
        "c_ten"             : "商品描述",
        "c_dvt"             : "计量单位",
        "c_luong"           : "申报数量",
        "c_dgia"            : "单价",
        "c_tgia"            : "总价值",
        "c_tk"              : "原始报关单号",
        "c_ref"             : "参考代码",
        "c_stat"            : "审计状态",
        "logout"            : "🚪 退出登录",
    },
}


# =============================================================================
# 3.  PAGE CONFIG & CSS
# =============================================================================

st.set_page_config(
    layout="wide",
    page_title="EXIM Reconciliation Pro",
    page_icon="🏢",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
#MainMenu, footer, header { visibility: hidden; }
html, body, [class*="st-"] {
    font-size: 1.05rem !important;
    font-family: 'Segoe UI', Tahoma, sans-serif;
}
h1 { font-size: 2.1rem !important; font-weight: 800 !important; color: #1E3A8A; }
h3 { font-size: 1.5rem !important; font-weight: 700 !important; color: #334155; }
h4 { font-size: 1.2rem !important; font-weight: 700 !important; color: #0F172A; }
.phase-box {
    border: 1px solid #E2E8F0;
    padding: 24px 28px;
    border-radius: 12px;
    background: linear-gradient(145deg, #F8FAFC, #FFFFFF);
    margin-bottom: 24px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.06);
}
.stButton > button {
    border-radius: 6px;
    font-weight: 700;
    letter-spacing: 0.4px;
    transition: all 0.18s ease;
}
.stButton > button:hover { transform: translateY(-2px); box-shadow: 0 4px 14px rgba(0,0,0,0.12); }
[data-testid="stMetric"] {
    background: #F1F5F9;
    border-radius: 8px;
    padding: 12px 16px;
    border-left: 4px solid #3B82F6;
}
[data-testid="stSidebar"] { background-color: #F8FAFC !important; }
</style>
""", unsafe_allow_html=True)


# =============================================================================
# 4.  LANGUAGE SELECTOR  (sidebar — runs before any widget that uses T)
# =============================================================================

with st.sidebar:
    sel_lang = st.selectbox("🌐 Language / Ngôn ngữ / 语言:", list(LANG.keys()), key="lang_sel")

T: Dict = LANG[sel_lang]


# =============================================================================
# 5.  AUTHENTICATION
# =============================================================================

def _auth_gate() -> bool:
    if st.session_state.get("auth"):
        return True
    _, mid, _ = st.columns([1, 2, 1])
    with mid:
        st.markdown(
            f"<h1 style='text-align:center;margin-top:8vh'>{T['gate_title']}</h1>",
            unsafe_allow_html=True,
        )
        st.markdown(
            f"<p style='text-align:center;color:#64748B;font-size:1.05em'>{T['gate_sub']}</p>",
            unsafe_allow_html=True,
        )
        pwd = st.text_input(T["pwd_label"], type="password", key="pwd_input")
        if pwd:
            expected = os.environ.get(
                "EXIM_PASSWORD",
                st.secrets.get("app_password", "ChingLuh@2026"),
            )
            if pwd == expected:
                st.session_state["auth"] = True
                log.info("Auth: login successful.")
                st.rerun()
            else:
                st.error(T["pwd_err"])
                log.warning("Auth: failed attempt.")
    return False

if not _auth_gate():
    st.stop()


# =============================================================================
# 6.  OCR ENGINE  (optional — gracefully degraded when easyocr not installed)
# =============================================================================

@st.cache_resource(show_spinner="⏳ Booting OCR engine…")
def _init_ocr():
    try:
        import easyocr
        return easyocr.Reader(["vi", "en"], gpu=False)
    except ImportError:
        return None

ocr = _init_ocr()


# =============================================================================
# 7.  DATA VALIDATOR
# =============================================================================

class DataValidator:
    @staticmethod
    def purify_code(val: Any) -> str:
        """Strip everything except A-Z 0-9; handle trailing '.0' artefacts."""
        if val is None or (isinstance(val, float) and np.isnan(val)):
            return ""
        s = str(val).strip()
        if s.endswith(".0"):
            s = s[:-2]
        return re.sub(r"[^A-Z0-9]", "", s.upper())

    @staticmethod
    def to_numeric(series) -> pd.Series:
        """Coerce any column to float — handles thousand-separators, stray text."""
        if series is None:
            return pd.Series(dtype=float)
        if isinstance(series, pd.DataFrame):
            series = series.iloc[:, 0]
        cleaned = (
            series.astype(str)
            .str.replace(",", "", regex=False)
            .str.replace(r"[^\d\-\.]", "", regex=True)
        )
        return pd.to_numeric(cleaned, errors="coerce").fillna(0.0)


# =============================================================================
# 8.  SMART FILE READER
# =============================================================================

class SmartReader:
    @staticmethod
    def _dedup(cols: list) -> list:
        seen: dict = {}
        out = []
        for c in cols:
            k = str(c).strip()
            if k in seen:
                seen[k] += 1
                out.append(f"{k}_{seen[k]}")
            else:
                seen[k] = 0
                out.append(k)
        return out

    @staticmethod
    def read(uf, keywords: list[str]) -> pd.DataFrame:
        if not uf:
            return pd.DataFrame()

        ext = uf.name.rsplit(".", 1)[-1].lower()
        log.info(f"Reading: {uf.name}  [{ext}]")

        try:
            # ── TABULAR ────────────────────────────────────────────────────
            if ext in {"xlsx", "xls", "xlsb", "xlsm", "csv", "txt"}:
                df: pd.DataFrame | None = None

                if ext in {"csv", "txt"}:
                    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
                        try:
                            uf.seek(0)
                            df = pd.read_csv(uf, on_bad_lines="skip", encoding=enc)
                            break
                        except Exception:
                            continue

                if df is None:
                    for eng in (None, "openpyxl", "xlrd"):
                        try:
                            uf.seek(0)
                            df = pd.read_excel(uf) if eng is None else pd.read_excel(uf, engine=eng)
                            if df is not None and not df.empty:
                                break
                        except Exception:
                            continue

                if df is None or df.empty:
                    log.warning(f"Empty or unreadable: {uf.name}")
                    return pd.DataFrame()

                df.columns = SmartReader._dedup(df.columns.tolist())

                # Auto-detect real header row (scan first 40 rows)
                hdr_idx: int | None = None
                for i, row in df.head(40).iterrows():
                    txt = " ".join(str(v).lower() for v in row.dropna())
                    if any(kw.lower() in txt for kw in keywords):
                        hdr_idx = i
                        break

                if hdr_idx is not None:
                    new_cols = df.iloc[hdr_idx].tolist()
                    df = df.iloc[hdr_idx + 1 :].copy()
                    df.columns = SmartReader._dedup(new_cols)
                    # Drop unnamed / nan columns and blank rows
                    df = df.loc[
                        :,
                        ~df.columns.str.contains(
                            r"^nan$|^Unnamed", case=False, regex=True
                        ),
                    ].dropna(how="all")
                    # Forward-fill first keyword column (merged cells)
                    for c in df.columns:
                        if any(kw.lower() in str(c).lower() for kw in keywords):
                            df[c] = df[c].ffill()
                            break

                log.info(f"  → {len(df)} rows parsed from {uf.name}")
                return df.reset_index(drop=True)

            # ── PDF ────────────────────────────────────────────────────────
            elif ext == "pdf":
                rows: list = []
                with pdfplumber.open(io.BytesIO(uf.read())) as pdf:
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        if tables:
                            for tbl in tables:
                                for row in tbl:
                                    clean = [
                                        str(c).strip().replace("\n", " ") if c else ""
                                        for c in row
                                    ]
                                    if any(clean):
                                        rows.append(clean)
                        else:
                            txt = page.extract_text()
                            if txt:
                                for line in txt.split("\n"):
                                    rows.append([line])
                if rows:
                    return pd.DataFrame(
                        rows[1:], columns=SmartReader._dedup(rows[0])
                    ).reset_index(drop=True)

            # ── IMAGE / OCR ────────────────────────────────────────────────
            elif ext in {"png", "jpg", "jpeg"}:
                if not ocr:
                    st.warning("EasyOCR not installed — cannot read image files.")
                    return pd.DataFrame()
                img = Image.open(uf).convert("RGB")
                img.thumbnail((1800, 1800), Image.Resampling.LANCZOS)
                texts = ocr.readtext(np.array(img), detail=0)
                return pd.DataFrame(texts, columns=["OCR_Text"])

        except Exception as exc:
            trace = traceback.format_exc()
            log.error(f"Parse failure: {uf.name}", trace)
            st.error(T["file_broken"].format(name=uf.name))

        return pd.DataFrame()


# =============================================================================
# 9.  RECONCILIATION ENGINE  (3-Way Gate + FIFO)
# =============================================================================

class ReconciliationEngine:

    def __init__(self, dfs: dict, ci: dict):
        self.dfs = dfs    # raw DataFrames keyed by name
        self.ci  = ci     # column-index overrides from sidebar
        self.V   = DataValidator

    # ── Column finder ────────────────────────────────────────────────────────
    def _col(self, df: pd.DataFrame, kws: list, fb: int) -> str:
        kws_l = [k.lower() for k in kws]
        for c in df.columns:
            if any(k in str(c).lower() for k in kws_l):
                return c
        idx = min(fb, len(df.columns) - 1)
        return df.columns[idx]

    # ── Build one ECUS output row ────────────────────────────────────────────
    def _row(
        self,
        hs: str, ma: str, desc: str, unit: str,
        qty: float, price: float, tk: str, status: str,
    ) -> dict:
        return {
            T["c_mshq"]  : hs,
            T["c_malieu"]: ma,
            T["c_ten"]   : desc,
            T["c_dvt"]   : unit,
            T["c_luong"] : round(qty, 2),
            T["c_dgia"]  : round(price, 4),
            T["c_tgia"]  : round(qty * price, 2),
            T["c_tk"]    : tk,
            T["c_ref"]   : ma,
            T["c_stat"]  : status,
        }

    # ── Status priority: Red > Yellow > Green ────────────────────────────────
    @staticmethod
    def _status(err_prefix: str, warn_prefix: str, ok_text: str, warn_sfx: str) -> str:
        """
        Compose a final status string with correct colour priority.
          - Any 🔴 prefix → red (error dominates)
          - Any 🟡 prefix → yellow (warning only when no red)
          - Otherwise     → green
        """
        if err_prefix:
            # Red error: append warning note if present, skip green
            return err_prefix + (warn_prefix.rstrip(" | ") if warn_prefix else "").strip()
        if warn_prefix:
            return warn_prefix + warn_sfx
        return ok_text

    # ── Main pipeline ────────────────────────────────────────────────────────
    def run(self) -> pd.DataFrame:
        ci = self.ci
        log.info(f"Pipeline start — {len(self.dfs)} source files.")

        # ── PKL demand dict ───────────────────────────────────────────────
        df_pkl = self.dfs["pkl"]
        pkl_mat = self._col(df_pkl, ["Material code", "Material", "Mã"], ci["inv_mat"])
        pkl_qty = self._col(df_pkl, ["Quantity", "PO Qty", "Q'TY", "Lượng"], ci["inv_qty"])
        _pk = df_pkl[[pkl_mat, pkl_qty]].copy()
        _pk["k"] = _pk[pkl_mat].apply(self.V.purify_code)
        _pk["q"] = self.V.to_numeric(_pk[pkl_qty])
        dict_pkl = _pk.groupby("k")["q"].sum().to_dict()

        # ── MB52 physical stock dict ──────────────────────────────────────
        df_mb = self.dfs["mb52"]
        mb_mat = self._col(df_mb, ["Material", "Mã"], ci["mb_mat"])
        mb_stk = self._col(df_mb, ["Unrestricted", "Tồn", "Stock"], ci["mb_stk"])
        _mb = df_mb[[mb_mat, mb_stk]].copy()
        _mb["k"] = _mb[mb_mat].apply(self.V.purify_code)
        _mb["q"] = self.V.to_numeric(_mb[mb_stk])
        dict_mb52 = _mb.groupby("k")["q"].sum().to_dict()

        # ── ZMM12 accounting dict ─────────────────────────────────────────
        df_zmm = self.dfs["zmm12"]
        zmm_mat = self._col(df_zmm, ["Material"], ci["zmm_mat"])
        zmm_qty = self._col(df_zmm, ["In Qty", "Qty"], ci["zmm_qty"])
        _zm = df_zmm[[zmm_mat, zmm_qty]].copy()
        _zm["k"] = _zm[zmm_mat].apply(self.V.purify_code)
        _zm["q"] = self.V.to_numeric(_zm[zmm_qty])
        dict_zmm12 = _zm.groupby("k")["q"].sum().to_dict()

        # ── IOP01 conversion table ────────────────────────────────────────
        df_iop = self.dfs["iop01"]
        iop_mat = self._col(df_iop, ["Material"], ci["iop_mat"])
        iop_rt  = self._col(df_iop, ["NLRate", "Tỷ lệ"], ci["iop_rt"])
        iop_ut  = self._col(df_iop, ["NLUnit", "ĐVT"], ci["iop_ut"])
        iop_hs  = self._col(df_iop, ["HSCode", "Mã HS", "HS"], ci["iop_hs"])
        _iop = df_iop[[iop_mat, iop_rt, iop_ut, iop_hs]].copy()
        _iop.columns = ["k", "rate", "unit", "hs"]
        _iop["k"]    = _iop["k"].apply(self.V.purify_code)
        _iop["rate"] = self.V.to_numeric(_iop["rate"])
        _iop = _iop.drop_duplicates("k").set_index("k")

        # ── HD03 ledger ───────────────────────────────────────────────────
        df_hd = self.dfs["hd03"]
        hd_mat = self._col(df_hd, ["Mã nguyên liệu", "Mã NL", "Material"], ci["hd_mat"])
        hd_tk  = self._col(df_hd, ["Số tờ khai", "TK"], ci["hd_tk"])
        hd_bal = self._col(df_hd, ["Lượng tồn", "Balance", "Còn lại"], ci["hd_bal"])
        hd_pri = self._col(df_hd, ["Đơn giá", "Price"], ci["hd_pri"])
        hd_dsc = self._col(df_hd, ["Tên hàng", "Description"], ci["hd_dsc"])
        _hd = df_hd[[hd_mat, hd_tk, hd_bal, hd_pri, hd_dsc]].copy()
        _hd.columns = ["k", "tk", "bal", "price", "desc"]
        _hd["k"]     = _hd["k"].apply(self.V.purify_code)
        _hd["bal"]   = self.V.to_numeric(_hd["bal"])
        _hd["price"] = self.V.to_numeric(_hd["price"])
        _hd = _hd[_hd["k"] != ""].reset_index(drop=True)

        # ── Invoice demand ────────────────────────────────────────────────
        df_inv = self.dfs["inv"]
        inv_mat = self._col(df_inv, ["Material code", "Material"], ci["inv_mat"])
        inv_qty = self._col(df_inv, ["Quantity", "PO Qty"], ci["inv_qty"])
        _inv = df_inv[[inv_mat, inv_qty]].copy()
        _inv.columns = ["k", "q"]
        _inv["k"] = _inv["k"].apply(self.V.purify_code)
        _inv["q"] = self.V.to_numeric(_inv["q"])
        demand = _inv[_inv["k"] != ""].groupby("k")["q"].sum().reset_index()

        log.info(f"Demand: {len(demand)} material codes.")

        # ── FIFO allocation loop ──────────────────────────────────────────
        results: list[dict] = []

        for _, dr in demand.iterrows():
            ma   = dr["k"]
            qty_inv = float(dr["q"])
            if qty_inv <= 0:
                continue

            # Conversion params
            iop_row = _iop.loc[ma] if ma in _iop.index else None
            rate    = float(iop_row["rate"]) if iop_row is not None else 1.0
            unit    = str(iop_row["unit"])  if iop_row is not None else "UNK"
            hs      = str(iop_row["hs"])    if iop_row is not None else "---"
            target  = round(qty_inv * rate, 4)

            # ── Gate 1: INV vs PKL documentary check ─────────────────────
            qty_pkl  = dict_pkl.get(ma, 0.0)
            err_pkl  = ""
            if abs(qty_inv - qty_pkl) > 0.1:
                err_pkl = T["st_err_pkl"].format(
                    inv=round(qty_inv, 2), pkl=round(qty_pkl, 2)
                )
                log.warning(f"{ma}: INV/PKL mismatch — INV={qty_inv}, PKL={qty_pkl}")

            # ── Gate 2: MB52 physical stock check ────────────────────────
            mb_stock = dict_mb52.get(ma, 0.0)
            if mb_stock < qty_inv:
                results.append(self._row(
                    hs, ma, "---", unit, target, 0.0, "---",
                    self._status(
                        err_pkl or T["st_err_mb52"].format(stock=round(mb_stock, 2)),
                        "", "", "",
                    ),
                ))
                log.error(f"{ma}: blocked — MB52 stock {mb_stock} < demand {qty_inv}")
                continue

            # ── Gate 3: ZMM12 accounting check (warn only) ───────────────
            zmm_stock = dict_zmm12.get(ma, 0.0)
            warn_zmm  = T["st_warn_zmm"] if zmm_stock < target else ""

            # ── HD03 FIFO split ───────────────────────────────────────────
            hd_sub = _hd[_hd["k"] == ma].copy()
            if hd_sub.empty:
                results.append(self._row(
                    hs, ma, "---", unit, target, 0.0, "---",
                    self._status(
                        err_pkl or T["st_err_hd03_miss"], "", "", "",
                    ),
                ))
                log.warning(f"{ma}: not found in HD03 ledger.")
                continue

            remaining = target
            for _, hd in hd_sub.iterrows():
                if remaining <= 0:
                    break
                avail = float(hd["bal"])
                if avail <= 0:
                    continue
                take      = min(remaining, avail)
                remaining = round(remaining - take, 4)

                status = self._status(err_pkl, warn_zmm, T["st_ok"], T["st_warn_chk"])
                results.append(self._row(
                    hs, ma, str(hd["desc"]), unit,
                    take, float(hd["price"]), str(hd["tk"]),
                    status,
                ))

            # Residual after HD03 exhausted
            if remaining > 1e-6:
                results.append(self._row(
                    hs, ma, "---", unit, remaining, 0.0, "---",
                    self._status(
                        err_pkl or T["st_err_hd03_empty"], warn_zmm, "", T["st_warn_chk"],
                    ),
                ))
                log.warning(f"{ma}: HD03 depleted — residual {remaining}")

        log.info(f"Pipeline complete: {len(results)} ECUS rows generated.")
        return pd.DataFrame(results)


# =============================================================================
# 10. EXPORT MANAGER
# =============================================================================

class ExportManager:

    @staticmethod
    def to_excel(df: pd.DataFrame) -> bytes:
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="ECUS_FORM")
            wb = writer.book
            ws = writer.sheets["ECUS_FORM"]

            # Formats
            hdr = wb.add_format({
                "bold": True, "bg_color": "#0F172A", "font_color": "#F8FAFC",
                "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True,
            })
            red  = wb.add_format({"bg_color": "#FEE2E2", "font_color": "#991B1B", "border": 1})
            yel  = wb.add_format({"bg_color": "#FEF9C3", "font_color": "#854D0E", "border": 1})
            grn  = wb.add_format({"bg_color": "#D1FAE5", "font_color": "#065F46", "border": 1})
            norm = wb.add_format({"border": 1, "valign": "vcenter"})
            num  = wb.add_format({"border": 1, "num_format": "#,##0.00", "valign": "vcenter"})
            num4 = wb.add_format({"border": 1, "num_format": "#,##0.0000", "valign": "vcenter"})

            num_cols = {T["c_luong"], T["c_tgia"]}
            num4_cols = {T["c_dgia"]}
            stat_idx  = len(df.columns) - 1

            ws.set_row(0, 32)
            for ci, cn in enumerate(df.columns):
                ws.write(0, ci, cn, hdr)
                w = min(max(df[cn].astype(str).map(len).max(), len(cn)) + 4, 50)
                ws.set_column(ci, ci, w)

            for ri, row_data in df.iterrows():
                sv = str(row_data.iloc[stat_idx])
                row_fmt = grn if "🟢" in sv else yel if "🟡" in sv else red
                for ci, (cn, val) in enumerate(zip(df.columns, row_data)):
                    safe = val if not (isinstance(val, float) and np.isnan(val)) else 0
                    if ci == stat_idx:
                        ws.write(ri + 1, ci, str(safe), row_fmt)
                    elif cn in num_cols:
                        ws.write_number(ri + 1, ci, float(safe) if safe else 0.0, num)
                    elif cn in num4_cols:
                        ws.write_number(ri + 1, ci, float(safe) if safe else 0.0, num4)
                    else:
                        ws.write(ri + 1, ci, "" if isinstance(safe, float) and np.isnan(safe) else safe, norm)

            ws.freeze_panes(1, 0)
            ws.autofilter(0, 0, len(df), len(df.columns) - 1)

        return buf.getvalue()

    @staticmethod
    def to_word(df: pd.DataFrame) -> bytes:
        doc = Document()
        title = doc.add_heading(T["btn_docx"].replace("📝 ", ""), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph(
            f"System: EXIM Reconciliation Pro  |  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in meta.runs:
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            run.font.size = Pt(10)

        doc.add_paragraph("─" * 70)

        sc = T["c_stat"]
        df_err = df[~df[sc].astype(str).str.contains("🟢", regex=False)]

        if df_err.empty:
            doc.add_paragraph("✅ No discrepancies detected. All records are valid.")
        else:
            n_r = df_err[sc].str.contains("🔴", regex=False).sum()
            n_y = df_err[sc].str.contains("🟡", regex=False).sum()
            doc.add_paragraph(
                f"⚠️  {len(df_err)} issues found — 🔴 {n_r} errors  |  🟡 {n_y} warnings"
            )
            cols = [T["c_malieu"], T["c_luong"], T["c_tk"], T["c_stat"]]
            cols = [c for c in cols if c in df.columns]

            tbl = doc.add_table(rows=1, cols=len(cols))
            tbl.style = "Table Grid"
            for i, h in enumerate(cols):
                cell = tbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True

            for _, r in df_err.head(250).iterrows():
                cells = tbl.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = str(r.get(c, ""))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# =============================================================================
# 11. SIDEBAR CONTROLS  (run after T is resolved)
# =============================================================================

with st.sidebar:
    st.markdown(f"**{T['cfg_tol']}**")
    tol_w = st.slider(T["tol_w"], 0.0, 1.0,  0.20, 0.05)
    tol_c = st.slider(T["tol_c"], 0.0, 0.10, 0.01, 0.01)
    st.markdown("---")

    with st.expander(T["cfg_col"], expanded=False):
        ci = {
            "inv_mat" : st.number_input(T["ci_inv_mat"],  value=1,  min_value=0, step=1),
            "inv_qty" : st.number_input(T["ci_inv_qty"],  value=5,  min_value=0, step=1),
            "mb_mat"  : st.number_input(T["ci_mb_mat"],   value=1,  min_value=0, step=1),
            "mb_stk"  : st.number_input(T["ci_mb_stk"],   value=3,  min_value=0, step=1),
            "iop_mat" : st.number_input(T["ci_iop_mat"],  value=1,  min_value=0, step=1),
            "iop_rt"  : st.number_input(T["ci_iop_rt"],   value=7,  min_value=0, step=1),
            "iop_ut"  : st.number_input(T["ci_iop_ut"],   value=6,  min_value=0, step=1),
            "iop_hs"  : st.number_input(T["ci_iop_hs"],   value=8,  min_value=0, step=1),
            "hd_mat"  : st.number_input(T["ci_hd_mat"],   value=2,  min_value=0, step=1),
            "hd_tk"   : st.number_input(T["ci_hd_tk"],    value=0,  min_value=0, step=1),
            "hd_bal"  : st.number_input(T["ci_hd_bal"],   value=4,  min_value=0, step=1),
            "hd_pri"  : st.number_input(T["ci_hd_pri"],   value=5,  min_value=0, step=1),
            "hd_dsc"  : st.number_input(T["ci_hd_dsc"],   value=3,  min_value=0, step=1),
            "zmm_mat" : st.number_input(T["ci_zmm_mat"],  value=10, min_value=0, step=1),
            "zmm_qty" : st.number_input(T["ci_zmm_qty"],  value=13, min_value=0, step=1),
        }

    st.markdown("---")
    if st.button(T["logout"], use_container_width=True):
        st.session_state.pop("auth", None)
        log.info("User logged out.")
        st.rerun()


# =============================================================================
# 12. MAIN HEADER
# =============================================================================

st.markdown(f"<h1>{T['main_title']}</h1>", unsafe_allow_html=True)
st.markdown(f"<p style='color:#475569;font-size:1.05em'>{T['main_sub']}</p>", unsafe_allow_html=True)
st.markdown("---")

tab_engine, tab_dash, tab_logs = st.tabs(
    [T["tab_engine"], T["tab_dashboard"], T["tab_logs"]]
)


# =============================================================================
# 13. TAB 1 — MAIN ENGINE
# =============================================================================

with tab_engine:

    # ── File upload lake ──────────────────────────────────────────────────────
    st.markdown("<div class='phase-box'>", unsafe_allow_html=True)
    st.markdown(f"#### {T['dl_title']}")
    st.caption(T["dl_sub"])

    FMT = ["xlsx", "csv", "xls", "xlsb", "xlsm", "txt", "pdf", "png", "jpg", "jpeg"]

    c1, c2, c3 = st.columns(3)
    with c1: f_inv   = st.file_uploader(T["f_inv"],   type=FMT, key="up_inv")
    with c2: f_pkl   = st.file_uploader(T["f_pkl"],   type=FMT, key="up_pkl")
    with c3: f_hd03  = st.file_uploader(T["f_hd03"],  type=FMT, key="up_hd03")

    c4, c5, c6 = st.columns(3)
    with c4: f_zmm12 = st.file_uploader(T["f_zmm12"], type=FMT, key="up_zmm12")
    with c5: f_iop01 = st.file_uploader(T["f_iop01"], type=FMT, key="up_iop01")
    with c6: f_mb52  = st.file_uploader(T["f_mb52"],  type=FMT, key="up_mb52")

    st.markdown("</div>", unsafe_allow_html=True)

    # ── Engine execution ──────────────────────────────────────────────────────
    if all([f_inv, f_pkl, f_hd03, f_zmm12, f_iop01, f_mb52]):

        st.markdown("<div class='phase-box'>", unsafe_allow_html=True)
        st.markdown(f"#### {T['eng_title']}")

        with st.spinner(T["spin_msg"]):
            raw = {
                "inv"  : SmartReader.read(f_inv,   ["Material", "Material code"]),
                "pkl"  : SmartReader.read(f_pkl,   ["Material", "Mã"]),
                "hd03" : SmartReader.read(f_hd03,  ["Mã nguyên liệu", "Material", "Mã NL"]),
                "iop01": SmartReader.read(f_iop01, ["Material", "NLClass"]),
                "mb52" : SmartReader.read(f_mb52,  ["Material", "Unrestricted"]),
                "zmm12": SmartReader.read(f_zmm12, ["PO Number", "Material"]),
            }

            broken = [k for k, v in raw.items() if v.empty]
            if broken:
                for k in broken:
                    st.error(T["file_broken"].format(name=k.upper()))
                    log.error(f"Abort: file {k} is empty after parsing.")
            else:
                engine = ReconciliationEngine(raw, ci)
                df_out = engine.run()
                st.session_state["ecus_output"] = df_out

        if "ecus_output" in st.session_state and not st.session_state["ecus_output"].empty:
            df_out = st.session_state["ecus_output"]
            st.success(T["succ_msg"])

            # ── Metric bar ────────────────────────────────────────────────────
            sc = T["c_stat"]
            n_ok   = int(df_out[sc].str.contains("🟢", regex=False).sum())
            n_warn = int(df_out[sc].str.contains("🟡", regex=False).sum())
            n_err  = int(df_out[sc].str.contains("🔴", regex=False).sum())
            tot_v  = float(df_out[T["c_tgia"]].sum())

            m1, m2, m3, m4, m5 = st.columns(5)
            m1.metric(T["m_total"], len(df_out))
            m2.metric(T["m_ok"],   n_ok)
            m3.metric(T["m_warn"], n_warn,
                      delta=f"-{n_warn}" if n_warn else None, delta_color="inverse")
            m4.metric(T["m_err"],  n_err,
                      delta=f"-{n_err}"  if n_err  else None, delta_color="inverse")
            m5.metric(T["m_value"], f"${tot_v:,.2f}")

            st.divider()

            # ── Filter + table ────────────────────────────────────────────────
            err_only = st.toggle(T["err_toggle"], value=False)
            disp = (
                df_out[~df_out[sc].str.contains("🟢", regex=False)]
                if err_only
                else df_out
            )

            def _style(v: str) -> str:
                s = str(v)
                if "🔴" in s: return "background:#fee2e2;color:#991b1b;font-weight:700"
                if "🟡" in s: return "background:#fef9c3;color:#854d0e;font-weight:700"
                if "🟢" in s: return "background:#d1fae5;color:#065f46;font-weight:700"
                return ""

            st.data_editor(
                disp.style
                    .map(_style, subset=[sc])
                    .format({
                        T["c_luong"]: "{:,.2f}",
                        T["c_dgia"] : "{:,.4f}",
                        T["c_tgia"] : "{:,.2f}",
                    }),
                use_container_width=True,
                hide_index=True,
                height=520,
            )

            # ── Downloads ─────────────────────────────────────────────────────
            st.write("")
            d1, d2 = st.columns(2)
            ts = datetime.now().strftime("%Y%m%d_%H%M")

            with d1:
                st.download_button(
                    T["btn_xlsx"],
                    data=ExportManager.to_excel(df_out),
                    file_name=f"ECUS_{ts}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    type="primary",
                )
            with d2:
                st.download_button(
                    T["btn_docx"],
                    data=ExportManager.to_word(df_out),
                    file_name=f"Report_{ts}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

        st.markdown("</div>", unsafe_allow_html=True)

    else:
        st.info(T["miss_files"])


# =============================================================================
# 14. TAB 2 — DASHBOARD
# =============================================================================

with tab_dash:
    st.markdown(f"### 📊 {T['dash_title']}")

    df_dash: pd.DataFrame = st.session_state.get("ecus_output", pd.DataFrame())

    if df_dash.empty:
        st.warning(T["dash_empty"])
    else:
        sc = T["c_stat"]
        counts = {
            "🟢 " + T["m_ok"]  : int(df_dash[sc].str.contains("🟢", regex=False).sum()),
            "🟡 " + T["m_warn"]: int(df_dash[sc].str.contains("🟡", regex=False).sum()),
            "🔴 " + T["m_err"] : int(df_dash[sc].str.contains("🔴", regex=False).sum()),
        }

        left, right = st.columns(2)

        with left:
            st.markdown(f"#### {T['dash_health']}")
            chart_df = pd.DataFrame(
                {"Count": list(counts.values())}, index=list(counts.keys())
            )
            st.bar_chart(chart_df, color=["#3B82F6"], use_container_width=True)

        with right:
            st.markdown(f"#### {T['dash_top5']}")
            top5 = (
                df_dash[[T["c_malieu"], T["c_tgia"]]]
                .groupby(T["c_malieu"], as_index=False)[T["c_tgia"]]
                .sum()
                .nlargest(5, T["c_tgia"])
                .set_index(T["c_malieu"])
            )
            st.bar_chart(top5, color=["#10B981"], use_container_width=True)

        st.divider()
        # Summary table
        st.markdown(f"**{T['m_total']}:** {len(df_dash)}  |  "
                    f"**{T['m_value']}:** ${df_dash[T['c_tgia']].sum():,.2f}")


# =============================================================================
# 15. TAB 3 — SYSTEM LOGS
# =============================================================================

with tab_logs:
    st.markdown(f"### {T['log_title']}")

    col_hdr, col_btn = st.columns([4, 1])
    with col_btn:
        if st.button(T["log_clear"], use_container_width=True):
            log.clear()
            st.rerun()

    all_logs = log.all()
    if not all_logs:
        st.info(T["log_empty"])
    else:
        ICON = {LogLevel.INFO: "ℹ️", LogLevel.WARNING: "⚠️", LogLevel.ERROR: "🚨"}
        for entry in reversed(all_logs):
            ts  = entry.timestamp.strftime("%H:%M:%S")
            ico = ICON.get(entry.level, "•")
            with st.expander(
                f"{ico} [{ts}]  {entry.level.value}  —  {entry.message}",
                expanded=(entry.level == LogLevel.ERROR),
            ):
                st.code(
                    entry.details if entry.details else T["log_no_detail"],
                    language="text",
                )
